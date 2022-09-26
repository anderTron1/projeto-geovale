# -*- coding: utf-8 -*-
"""
Created on Wed Oct 20 14:41:31 2021

@author: André luiz Pires Guimarães
"""
from elementsAdditional import ElementsAdditional
from Layouts.keys_names.keys_to_registration import *
from Layouts.search_register_person import Search_register_person
from Layouts.import_contract import Import_contract
import PySimpleGUI as sg
import os
import time

import re

from decimal import Decimal
import locale
locale.setlocale(locale.LC_ALL, 'pt_BR.utf8')

from docx import Document

DEFAULT_KEY_INPUT_ID = '-ID_DO_REGISTRO_DA_BASE_DE_DADOS-'
DEFAULT_KEY_INPUT_NAME = '-NOME-'
DEFAULT_KEY_INPUT_CPF = '-CPF-'
DEFAULT_KEY_INPUT_CONTRACT = '-CONTRATO-'
DEFAULT_KEY_INPUT_NAME_FILE_TO_SAVE = '<<-NOME_DO_ARQUIVO_PARA_SARVAR->>'

DEFAULT_KEY_BTN_SEARCH_RECORD = '-PESQUISAR_REGISTRO-'
DEFAULT_KEY_BTN_SEARCH_CONTRACT = '-PESQUISAR_CONTRATO-'
DEFAULT_KEY_BTN_GENERATE = '-GERAR_CONTRATO-'
DEFAULT_KEY_BTN_CANCEL = '-CANCELAR_INFORMACOES-'
DEFAULT_KEY_PROGRESS_BAR_GENERATE_CONSTRUCT = 'PROGRESS_BAR_GENERATE_CONSTRUCT'
DEFAULT_KEY_PROCESSTO_CREATE_A_CONTRACT = 'PROCESSO_PARA_GERAR_CONTRATO'

class Generate:
    def __init__(self, name):
        self.document = Document(name)
        self.elements = ElementsAdditional()
        
    def lines(self,inline, dict_sub):
        
        for i in range(len(inline)):
            paragr = inline[i].text
            #print(paragr)
            for text in dict_sub.keys():
                if paragr.find(text) != -1:
                    
                    #print(paragr)
                    texto = paragr.replace(text, str(dict_sub[text]))
                    inline[i].text = texto
                    paragr = texto
    #def lines1(self,inline, dict_sub):
        
    '''
    Clear document keys that have no values
    '''
    def clear_keys_docx(self, keys):
        for paragraph in self.document.paragraphs:
            for key in keys:
                for textParag in paragraph.runs:
                    if key in textParag.text:
                    	textParag.text = textParag.text.replace(key,' ')
                    	print('limpando chave: ', key, ' da qual não existe informação disponivel neste registro')
    
    def update_docx(self, name, datas=None, is_table=False, progress_bar=None, income = 0):

    	if name.find('.docx') == -1:
        	sg.popup_error('O arquivo não é um formato docx')
    	else:
            if is_table == False:
                #update  paragraph
                #Verificar texto...
                for paragraph in self.document.paragraphs:
                    for key in datas.keys():
                        for textParag in paragraph.runs:
                            if key in textParag.text and str(datas[key]) != '--':
                                	textParag.text = textParag.text.replace(
                                	    key, str(datas[key]))
                                	#break

                #Verificar tabelas...
                for table in self.document.tables:
                    for line in table.rows:
                        for cell in line.cells:
                            for textParag in cell.paragraphs:
                                for key in datas.keys():
                                    for textTable in textParag.runs:
                                        if key in textTable.text and str(datas[key]) != '--':
                                            textTable.text = textTable.text.replace(
                                                key, str(datas[key]))
                return income
            if is_table == True and progress_bar != None:
                cont = 0
                size_data = len(datas)
                progress_bar.UpdateBar(0, size_data+1)
                time.sleep(.01)
                #update table
                for table in self.document.tables:
                    cont_rows = len(datas) / (len(table.columns)-1)

                    if table.rows[0].cells[0].text == DEFAULT_KEY_TABLE_RESIDENTS:
                        table.rows[0].cells[0].text = 'Cadastro de Moradores'

                        if len(table.rows) < cont_rows:
                            difference = cont_rows - len(table.rows)

                            for div in range(difference):
                                table.add_row().cells

                        for line in table.rows:
                            for cell in line.cells:
                                if cont == size_data:
                                    break
                                if cell.text == '':
                                    if str(datas[cont]).find('R$') != -1:
                                        income += float(re.sub('([^0-9].[^0-9]{1,2})', '', str(
                                            datas[cont]).replace('.', '').replace(',', '.')))
                                        cell.text = datas[cont]
                                    else:
                                        cell.text = str(datas[cont])
                                    cont += 1

                                    progress_bar.UpdateBar(cont, size_data)
                                    time.sleep(.01)
                #print(income)
                return income
                
                                   
    def save(self,dict_save, save_as): 
        save = True
        try:
            self.document.save(dict_save+'/'+save_as+'.docx')
        except:
            save = False
        return save
    
    def close(self):
        self.document.close()
        
        
class Generate_contract:
    def __init__(self, conectionDB):
        self._conn = conectionDB
        self._class_search_register_person = Search_register_person(conectionDB)
        self._class_import_contract = Import_contract(return_only_the_path=True)
        
        self.elemAdditional = ElementsAdditional()
        
        self._path_contract = None
        self._id_register_db = None

    def layout(self):
        construct_layout = [
                        [sg.T('ID:', size=(5)), sg.Input(size=(10),key=DEFAULT_KEY_INPUT_ID, disabled=True)],
                        [sg.T('Nome:', size=(5)), sg.Input(key=DEFAULT_KEY_INPUT_NAME, disabled=True)],
                        [sg.T('CPF:', size=(5)), sg.Input(size=(15),key=DEFAULT_KEY_INPUT_CPF, disabled=True)],
                        [sg.Button('Pesquisar', key=DEFAULT_KEY_BTN_SEARCH_RECORD)],
                        [sg.HorizontalSeparator()],
                        [sg.T('Contrato:'), sg.Input(key=DEFAULT_KEY_INPUT_CONTRACT, disabled=True)],
                        [sg.Button('Pesquisar', key=DEFAULT_KEY_BTN_SEARCH_CONTRACT, disabled=True)],
                        [sg.T('Salvar como:'), sg.Input(size=30, key=DEFAULT_KEY_INPUT_NAME_FILE_TO_SAVE, disabled=True), sg.T('.docx')],
                        [sg.Button('Gerar', key=DEFAULT_KEY_BTN_GENERATE, disabled=True),
                         sg.Button('Cancelar', key=DEFAULT_KEY_BTN_CANCEL, disabled=True)],
                        [sg.ProgressBar(1, orientation='h', size=(40,20), key=DEFAULT_KEY_PROGRESS_BAR_GENERATE_CONSTRUCT)],
                        [sg.T('', key=DEFAULT_KEY_PROCESSTO_CREATE_A_CONTRACT)]
                    ]
        return construct_layout
    
    def __database_records(self, window):
        KEY_ID = 0
        KEY_NAME = 1
        KEY_CPF = 2
        
        search = Search_register_person(self._conn)
        _, self._id_register_db = search.window_button_search()
            
        register_database = self._conn.select_register([self._conn.id_register_people,'name', 'cpf'], self._conn.register_people, self._conn.id_register_people, self._id_register_db)
        
        if len(register_database) > 0:
            window.Element(DEFAULT_KEY_INPUT_ID).update(register_database[0][KEY_ID])
            window.Element(DEFAULT_KEY_INPUT_NAME).update(register_database[0][KEY_NAME])
            window.Element(DEFAULT_KEY_INPUT_CPF).update(self.elemAdditional.valid_cpf(str(register_database[0][KEY_CPF])))
        
    def __contract_files(self, window, value):
        path = self._class_import_contract.exec_class()
        
        if path != -1:
            self._path_contract = path
            window.Element(DEFAULT_KEY_INPUT_CONTRACT).update(self._path_contract[self._path_contract.rindex('/')+1:len(self._path_contract)])
            window.Element(DEFAULT_KEY_INPUT_NAME_FILE_TO_SAVE).update(value[DEFAULT_KEY_INPUT_NAME])
            
    def __enabled_btns(self, window, disabled):
        window.Element(DEFAULT_KEY_BTN_GENERATE).update(disabled=disabled)
        window.Element(DEFAULT_KEY_BTN_CANCEL).update(disabled=disabled)
        
    def __clear_inputs(self, window, is_to_clean = True):
        window.Element(DEFAULT_KEY_INPUT_ID).update('')
        window.Element(DEFAULT_KEY_INPUT_NAME).update('')
        window.Element(DEFAULT_KEY_INPUT_CPF).update('')
        window.Element(DEFAULT_KEY_INPUT_CONTRACT).update('')
        window.Element(DEFAULT_KEY_INPUT_NAME_FILE_TO_SAVE).update('')
    
    def get_incomes(self, id_register):
        sql = self._conn.register_people + ' WHERE ' + self._conn.id_register_people + ' = ' + str(id_register)
        sql_spouse = self._conn.register_spouse + ' WHERE ' + self._conn.name_id_to_table_register + ' = ' + str(id_register)
            
        income_people = self._conn.select_register(['income_between'], sql)
        income_people = float(re.sub('([^0-9].[^0-9]{1,2})', '', income_people[0][0].replace('.', '').replace(',','.')))
            
        try:
            income_spouse = self._conn.select_register(['income_spouse'], sql_spouse)
            income_spouse = float(re.sub('([^0-9].[^0-9]{1,2})', '', income_spouse[0][0].replace('.', '').replace(',','.')))
        except:
            income_spouse = 0
        
        return income_people, income_spouse
    
    def get_db(self,window, value,keys_fields, name_register, name_id_register, id_register, is_table=False):
        fileds_and_field_db = self._conn._take_fields_records(name_register, keys_fields)
        progress_bar = window.Element(DEFAULT_KEY_PROGRESS_BAR_GENERATE_CONSTRUCT)
        
        if is_table:
            values = []            
            register_db = self._conn.select_register(fileds_and_field_db.keys(), name_register,name_id_register, id_register)
            if register_db != None:
                for val_tuple in register_db:
                    for value in val_tuple:
                        values.append(value)
                        
            new_income = self.__generate.update_docx(self._path_contract, values, is_table, progress_bar)
            #print('dadosss: ', new_income)
            return new_income
        else:
            key_value = dict()
            register_db = self._conn.select_register(fileds_and_field_db.keys(), name_register,name_id_register, id_register)[0]
        
            #print(register_db, '\n\n')
            size_list = len(fileds_and_field_db.values())
            progress_bar.UpdateBar(0, size_list)
            for cont, key in enumerate(fileds_and_field_db.values()):
                if register_db[cont] != None:
                    key_value[key] = register_db[cont]
            
                self.__generate.update_docx(self._path_contract, key_value, is_table)
                progress_bar.UpdateBar(cont, size_list)
                time.sleep(.01)
            #print(key_value)
            
    def get_datas_db(self,window, value):
        
        self.__generate = Generate(self._path_contract)
        
        save_in_directory = sg.tk.filedialog.askdirectory(initialdir=os.path.abspath(os.sep))
        
        #everyting register client
        name_register = self._conn.register_people
        name_id_register = self._conn.id_register_people
        id_register = int(value[DEFAULT_KEY_INPUT_ID])

        window.Element(DEFAULT_KEY_PROCESSTO_CREATE_A_CONTRACT).update('Inserindo registros dos dados pessoais...')
        self.get_db(window,value, keys_fields(), name_register, name_id_register, id_register)
        self.__generate.clear_keys_docx(keys_fields())
        
        #everyting register spouse
        name_spouse = self._conn.register_spouse
        name_id_to_register = self._conn.name_id_to_table_register
        
        register_spouse_exist = self._conn.query_record(name_spouse, name_id_to_register, id_register)
        
        if register_spouse_exist:
            window.Element(DEFAULT_KEY_PROCESSTO_CREATE_A_CONTRACT).update('Inserindo registros do cônjuge...')
            self.get_db(window,value, keys_fields_spouse(), name_spouse, name_id_to_register, id_register)
        self.__generate.clear_keys_docx(keys_fields_spouse())
            
        #everyting register residents
        name_register_residents = self._conn.register_residents
        
        register_residents_exist = self._conn.query_record(name_register_residents, name_id_to_register, id_register)
        
        income_total = 0
        if register_residents_exist:
            
            window.Element(DEFAULT_KEY_PROCESSTO_CREATE_A_CONTRACT).update('Inserindo registros nas tabelas...')
            income_total = self.get_db(window,value, key_fields_residents().pop(0), name_register_residents, name_id_to_register, id_register, is_table=True)
            
        try:
            income_people, income_spouse = self.get_incomes(id_register)
            #print('renda pessoa: ',income_people, 'renda spouse: ', income_spouse, 'renda total: ', income_total)
            income_total += income_people
            income_total += income_spouse
            income_total = locale.currency(income_total, grouping=True)
        except:
            sg.popup_error('ERRO para converter o valor da renda familiar')
                
        #update icome table
        #print(income_total)
        dic_icome = {DEFAULT_KEY_TXT_FAMILY_INCOME_REGIST_RESID: income_total}
        for paragraph in self.__generate.document.paragraphs:
            self.__generate.lines(paragraph.runs, dic_icome)
        
        self.__generate.clear_keys_docx(key_fields_residents())
        self.__generate.clear_keys_docx([DEFAULT_KEY_TXT_FAMILY_INCOME_REGIST_RESID])
            
        #generate and save changes to the Word file 
        save = self.__generate.save(save_in_directory, value[DEFAULT_KEY_INPUT_NAME_FILE_TO_SAVE])
        if save:
            window.Element(DEFAULT_KEY_PROCESSTO_CREATE_A_CONTRACT).update('Arquivo gerado e salvo com sucesso!...',text_color='white')
        else:
            window.Element(DEFAULT_KEY_PROCESSTO_CREATE_A_CONTRACT).update('Um erro inesperado foi encontrato ao tentar gerar e salvar o arquivo!...',text_color='red')
        
    def exec_class(self):
        window = sg.Window('Gerar contrato para registros', self.layout(),icon=r'image/iconLogo.ico', keep_on_top=False, modal=True)
        while(True):
            event, value = window.read(timeout=100)
            
            if event == sg.WINDOW_CLOSED:
                break
            
            if event == DEFAULT_KEY_BTN_SEARCH_RECORD:
                self.__database_records(window)
                
            if event == DEFAULT_KEY_BTN_SEARCH_CONTRACT:
                self.__contract_files(window, value)
                
            if value[DEFAULT_KEY_INPUT_NAME] == '':
                window.Element(DEFAULT_KEY_BTN_SEARCH_CONTRACT).update(disabled=True)
                window.Element(DEFAULT_KEY_INPUT_NAME_FILE_TO_SAVE).update(disabled=True)
            else:
                window.Element(DEFAULT_KEY_BTN_SEARCH_CONTRACT).update(disabled=False)
                window.Element(DEFAULT_KEY_INPUT_NAME_FILE_TO_SAVE).update(disabled=False)
                
            if value[DEFAULT_KEY_INPUT_NAME] != '' and value[DEFAULT_KEY_INPUT_CONTRACT] != '' and value[DEFAULT_KEY_INPUT_NAME_FILE_TO_SAVE] != '':
                self.__enabled_btns(window, False)
            else:
                self.__enabled_btns(window, True)
            
            if event == DEFAULT_KEY_BTN_GENERATE:
                self.get_datas_db(window, value)
                self.__clear_inputs(window, True)
                
            if event == DEFAULT_KEY_BTN_CANCEL:
                self.__clear_inputs(window, True)
            
        window.close()
